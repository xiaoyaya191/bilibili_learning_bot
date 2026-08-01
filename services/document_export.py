"""Markdown 知识笔记导出为 PDF / Word。"""
from __future__ import annotations

import asyncio
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from colorama import Fore, Style
except Exception:  # 非 CLI 场景（如网页端导入）降级为无色
    class _N:
        def __getattr__(self, _):
            return ""
    Fore = Style = _N()


BASE_DIR = Path(__file__).resolve().parent.parent
from core.config import CONFIG_FILE
from core.user_data import WORD_DIR



def _read_doc_cfg() -> dict:
    try:
        cfg_path = Path(CONFIG_FILE)
        if cfg_path.exists():
            with open(cfg_path, "r", encoding="utf-8") as f:
                return json.load(f).get("document_export", {}) or {}
    except Exception:
        pass
    return {}


def _resolve_out_dir(out_dir=None) -> Path:
    """解析 Word/PDF 导出目录：优先参数，其次配置中的 document_export.output_dir/folder_name，回退默认。"""
    if out_dir:
        p = Path(out_dir)
        return p if p.is_absolute() else (WORD_DIR if getattr(sys, "frozen", False) else BASE_DIR / p)

    cfg = _read_doc_cfg()
    d = cfg.get("output_dir") or cfg.get("folder_name")
    if not d:
        return WORD_DIR
    p = Path(d)
    return p if p.is_absolute() else (WORD_DIR if getattr(sys, "frozen", False) else BASE_DIR / p)


def _doc_prompt() -> str:
    return (_read_doc_cfg().get("prompt") or "").strip()


def _safe_name(path: Path) -> str:
    return re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", path.stem).strip("_") or "note"


def _md_to_plain_lines(text: str) -> list[str]:
    text = re.sub(r"```.*?```", "", text or "", flags=re.S)
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    lines = []
    for raw in text.splitlines():
        line = raw.strip()
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^[-*+]\s+", "• ", line)
        line = line.replace("**", "").replace("__", "").replace("`", "")
        lines.append(line)
    return lines


def _read_md(md_path: str | Path, kb_root: str | Path | None = None) -> Path:
    from core.config import config, resolve_knowledge_base_dir
    root = Path(kb_root or resolve_knowledge_base_dir(config)).resolve()

    target = Path(md_path)
    if not target.is_absolute():
        target = root / target
    target = target.resolve()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise ValueError("禁止导出知识库外部文件") from exc
    if not target.exists() or not target.is_file():
        raise FileNotFoundError("知识文件不存在")
    return target


def export_docx(md_path: str | Path, kb_root: str | Path | None = None, out_dir: str | Path | None = None) -> str:
    target = _read_md(md_path, kb_root)
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先安装依赖：pip install python-docx") from exc
    text = target.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    doc.add_heading(target.stem, level=1)
    for line in _md_to_plain_lines(text):
        if not line:
            doc.add_paragraph("")
        elif line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    output_dir = _resolve_out_dir(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    prompt = _doc_prompt()
    if prompt:
        ppre = doc.add_paragraph()
        r = ppre.add_run(prompt)
        r.italic = True
        doc.add_paragraph("")
    out = output_dir / f"{_safe_name(target)}.docx"
    doc.save(str(out))
    return str(out)


def export_pdf(md_path: str | Path, kb_root: str | Path | None = None, out_dir: str | Path | None = None) -> str:
    target = _read_md(md_path, kb_root)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("缺少 reportlab，请先安装依赖：pip install reportlab") from exc

    output_dir = _resolve_out_dir(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    out = output_dir / f"{_safe_name(target)}.pdf"
    c = canvas.Canvas(str(out), pagesize=A4)
    width, height = A4
    font_name = "Helvetica"
    for candidate in ["C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf", "C:/Windows/Fonts/simsun.ttc"]:
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("CNFont", candidate))
                font_name = "CNFont"
                break
            except Exception:
                pass
    c.setFont(font_name, 16)
    y = height - 48
    c.drawString(42, y, target.stem[:60])
    y -= 32
    c.setFont(font_name, 10)
    for line in _md_to_plain_lines(target.read_text(encoding="utf-8", errors="replace")):
        chunks = [line[i:i + 52] for i in range(0, len(line), 52)] or [""]
        for chunk in chunks:
            if y < 42:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 42
            c.drawString(42, y, chunk)
            y -= 15
    c.save()
    return str(out)


def export_document(md_path: str | Path, fmt: str, kb_root: str | Path | None = None, out_dir: str | Path | None = None) -> str:
    fmt = (fmt or "").lower().strip()
    if fmt in {"docx", "word"}:
        return export_docx(md_path, kb_root=kb_root, out_dir=out_dir)
    if fmt == "pdf":
        return export_pdf(md_path, kb_root=kb_root, out_dir=out_dir)
    raise ValueError("仅支持 pdf/docx")


# ─────────────────────────────────────────────
#  从内存文本直接导出（绕过知识库目录限制，供视频内容等任意文本使用）
# ─────────────────────────────────────────────
def _safe_title(title: str) -> str:
    return re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", (title or "note")).strip("_") or "note"


def export_docx_text(text: str, title: str, out_dir: str | Path | None = None) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先安装依赖：pip install python-docx") from exc
    doc = Document()
    doc.add_heading(title or "笔记", level=1)
    for line in _md_to_plain_lines(text or ""):
        if not line:
            doc.add_paragraph("")
        elif line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    output_dir = _resolve_out_dir(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    prompt = _doc_prompt()
    if prompt:
        ppre = doc.add_paragraph()
        r = ppre.add_run(prompt)
        r.italic = True
        doc.add_paragraph("")
    out = output_dir / f"{_safe_title(title)}.docx"
    doc.save(str(out))
    return str(out)


def export_pdf_text(text: str, title: str, out_dir: str | Path | None = None) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("缺少 reportlab，请先安装依赖：pip install reportlab") from exc
    output_dir = _resolve_out_dir(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    out = output_dir / f"{_safe_title(title)}.pdf"
    c = canvas.Canvas(str(out), pagesize=A4)
    width, height = A4
    font_name = "Helvetica"
    for candidate in ["C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf", "C:/Windows/Fonts/simsun.ttc"]:
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("CNFont", candidate))
                font_name = "CNFont"
                break
            except Exception:
                pass
    c.setFont(font_name, 16)
    y = height - 48
    c.drawString(42, y, (title or "note")[:60])
    y -= 32
    c.setFont(font_name, 10)
    for line in _md_to_plain_lines(text or ""):
        chunks = [line[i:i + 52] for i in range(0, len(line), 52)] or [""]
        for chunk in chunks:
            if y < 42:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 42
            c.drawString(42, y, chunk)
            y -= 15
    c.save()
    return str(out)


# ─────────────────────────────────────────────
#  CLI 交互：把「同一份视频内容」导出为 Word / PDF / PPT
#  W 命令（视频→网页）与 V 命令（手动视频分析）共用，避免重复代码
# ─────────────────────────────────────────────
async def export_video_content(title: str, up_name: str, video_url: str, ctx: str,
                                formats: list[str], stats: dict | None = None, desc: str = "",
                                bvid: str | None = None, brain=None, ppt_theme: str = "dark",
                                ppt_detail: str = "medium") -> dict:
    """非交互：把视频内容导出为指定格式列表（formats ∈ {'docx','pdf','ppt','mindmap'}）。
    ppt_detail: 'simple' 简单 | 'medium' 中长(默认) | 'detailed' 详细
    返回 {fmt: {'path': str} | {'error': str}}，供 CLI 与 Web 共用。"""
    results: dict = {}
    _st = stats or {}
    note = f"# {title}\n\n> UP主: {up_name} · 链接: {video_url}\n"
    note += (f"> 播放:{_st.get('view', '?')} 点赞:{_st.get('like', '?')} 投币:{_st.get('coin', '?')} "
             f"收藏:{_st.get('favorite', '?')} 弹幕:{_st.get('danmaku', '?')} "
             f"评论:{_st.get('comment', '?')} 时长:{_st.get('duration', '?')}\n\n")
    if desc:
        note += f"**视频简介**：{desc}\n\n"
    note += f"## 内容\n\n{ctx}\n"

    fmt_set = {str(f).lower() for f in (formats or [])}
    # Word
    if 'docx' in fmt_set or 'word' in fmt_set:
        try:
            results['docx'] = {'path': export_docx_text(note, title)}
        except Exception as _e:
            results['docx'] = {'error': str(_e)}
    # PDF
    if 'pdf' in fmt_set:
        try:
            results['pdf'] = {'path': export_pdf_text(note, title)}
        except Exception as _e:
            results['pdf'] = {'error': str(_e)}
    # PPT（复用既有 generate_ppt_from_bvid）
    if 'ppt' in fmt_set:
        try:
            from services._services_ai import _live_config
            from services.video_to_ppt import generate_ppt_from_bvid
            live = _live_config()
            ppt = await generate_ppt_from_bvid(
                bvid=bvid, api_key=live.get('api_key', ''), base_url=live.get('base_url', ''),
                model=live.get('model_brain', ''), cookies_obj=getattr(brain, 'cookies', None),
                theme=ppt_theme, detail_level=ppt_detail, open_browser=False, auto_save=True)
            if ppt.get('success'):
                results['ppt'] = {'path': ppt.get('html_path')}
            else:
                results['ppt'] = {'error': ppt.get('error') or 'PPT 生成失败'}
        except Exception as _e:
            results['ppt'] = {'error': str(_e)}
    # 思维导图（从文本生成临时 .md，调用 export_mindmap）
    if 'mindmap' in fmt_set:
        try:
            import tempfile
            from services.mindmap_export import export_mindmap
            _safe = re.sub(r'[\\/:*?"<>|]', '_', title)[:80]
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', prefix=f'[{bvid or "video"}] - ',
                                             encoding='utf-8', delete=False) as _tf:
                _tf.write(note)
                _tmp_path = _tf.name
            try:
                mm_path = export_mindmap(_tmp_path)
                results['mindmap'] = {'path': mm_path}
            finally:
                Path(_tmp_path).unlink(missing_ok=True)
        except Exception as _e:
            results['mindmap'] = {'error': str(_e)}
    return results


async def export_video_content_interactive(title: str, up_name: str, video_url: str, ctx: str,
                                            stats: dict | None = None, desc: str = "",
                                            bvid: str | None = None, brain=None):
    """CLI 交互版：提示用户选择格式后调用 export_video_content。W/V 命令共用。"""
    try:
        print(f"\n{Fore.CYAN}📦 是否同时把该视频内容导出为其他格式？(基于同一份视频内容){Style.RESET_ALL}")
        print(f"  {Fore.YELLOW}1.{Style.RESET_ALL} 📄 Word 文档 (.docx)")
        print(f"  {Fore.YELLOW}2.{Style.RESET_ALL} 📑 PDF 文档 (.pdf)")
        print(f"  {Fore.YELLOW}3.{Style.RESET_ALL} 🎞️ PPT 演示 (.html)")
        print(f"  {Fore.YELLOW}4.{Style.RESET_ALL} 🧠 思维导图 (.html)")
        print(f"  {Fore.CYAN}(可多选，如 1234 / 1 / 直接回车跳过){Style.RESET_ALL}")
        loop = asyncio.get_running_loop()
        fmt_choice = (await loop.run_in_executor(None, input, f"{Fore.GREEN}> {Style.RESET_ALL}")).strip()
        if not fmt_choice:
            return
        fm: list[str] = []
        if '1' in fmt_choice:
            fm.append('docx')
        if '2' in fmt_choice:
            fm.append('pdf')
        if '3' in fmt_choice:
            fm.append('ppt')
        if '4' in fmt_choice:
            fm.append('mindmap')
        if not fm:
            return
        # PPT 风格选择
        ppt_theme = "dark"
        ppt_detail = "medium"
        if 'ppt' in fm:
            print(f"\n{Fore.CYAN}🎨 PPT 视觉风格:{Style.RESET_ALL}")
            print(f"  {Fore.YELLOW}1.{Style.RESET_ALL} dark — 深色科技风 (默认)")
            print(f"  {Fore.YELLOW}2.{Style.RESET_ALL} claude — Claude暖橙风格 (Fraunces+Inter)")
            print(f"  {Fore.YELLOW}3.{Style.RESET_ALL} claude_slides — Claude幻灯片 (纯白+暖橙+翻页)")
            print(f"  {Fore.YELLOW}4.{Style.RESET_ALL} claude_slides_v2 — Claude动画幻灯片v2 (11种动画)")
            print(f"  {Fore.YELLOW}5.{Style.RESET_ALL} purple — 紫色主题")
            print(f"  {Fore.YELLOW}6.{Style.RESET_ALL} cyan — 青色主题")
            theme_choice = (await loop.run_in_executor(None, input, f"{Fore.GREEN}  PPT风格 (1-6, 回车=dark): {Style.RESET_ALL}")).strip()
            theme_map = {'1': 'dark', '2': 'claude', '3': 'claude_slides', '4': 'claude_slides_v2', '5': 'purple', '6': 'cyan'}
            ppt_theme = theme_map.get(theme_choice, 'dark')
            print(f"  {Fore.GREEN}[OK] PPT风格: {ppt_theme}{Style.RESET_ALL}")

            print(f"\n{Fore.CYAN}📊 PPT 内容详略:{Style.RESET_ALL}")
            print(f"  {Fore.YELLOW}1.{Style.RESET_ALL} 简单 — 核心要点，3-5个观点，快速浏览")
            print(f"  {Fore.YELLOW}2.{Style.RESET_ALL} 中长 — 适中展开 (默认)")
            print(f"  {Fore.YELLOW}3.{Style.RESET_ALL} 详细 — 完整深入，包含例子和细节")
            detail_choice = (await loop.run_in_executor(None, input, f"{Fore.GREEN}  详略级别 (1-3, 回车=中长): {Style.RESET_ALL}")).strip()
            detail_map = {'1': 'simple', '2': 'medium', '3': 'detailed'}
            ppt_detail = detail_map.get(detail_choice, 'medium')
            print(f"  {Fore.GREEN}[OK] 详略级别: {detail_map.get(detail_choice, '中长')}{Style.RESET_ALL}")
        res = await export_video_content(title, up_name, video_url, ctx, fm,
                                         stats=stats, desc=desc, bvid=bvid, brain=brain,
                                         ppt_theme=ppt_theme, ppt_detail=ppt_detail)
        for _f, _r in res.items():
            if 'path' in _r:
                print(f"{Fore.GREEN}  ✅ {_f.upper()} 已导出: {_r['path']}{Style.RESET_ALL}")
            else:
                print(f"{Fore.RED}  ⚠ {_f.upper()} 导出失败: {_r['error']}{Style.RESET_ALL}")
    except Exception as _e:
        print(f"{Fore.YELLOW}  ⚠ 附加导出已跳过: {_e}{Style.RESET_ALL}")
